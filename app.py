import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores import Chroma
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from question_paper_template import create_question_paper_pdf
from answer_sheet_template import create_answer_sheet_pdf
import tempfile

load_dotenv()
os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))






def get_pdf_text(pdf_docs):
    text=""
    for pdf in pdf_docs:
        pdf_reader= PdfReader(pdf)
        for page in pdf_reader.pages:
            text+= page.extract_text()
    return  text



def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks


def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
    vector_store = Chroma.from_texts(text_chunks, embedding=embeddings, persist_directory="chroma_db")
    vector_store.persist()


def get_conversational_chain():

    prompt_template = """
    Answer the question as detailed as possible from the provided context, make sure to provide all the details, if the answer is not in
    provided context just say, "answer is not available in the context", don't provide the wrong answer\n\n
    Context:\n {context}?\n
    Question: \n{question}\n

    Answer:
    """

    model = ChatGoogleGenerativeAI(model="gemini-2.0-flash",
                             temperature=0.3)

    prompt = PromptTemplate(template = prompt_template, input_variables = ["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

    return chain



def user_input(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
    new_db = Chroma(persist_directory="chroma_db", embedding_function=embeddings)
    docs = new_db.similarity_search(user_question)
    chain = get_conversational_chain()
    response = chain(
        {"input_documents":docs, "question": user_question}
        , return_only_outputs=True)
    print(response)
    st.write("Reply: ", response["output_text"])

def generate_questions_from_pdf(pdf_docs, num_questions, output_format, difficulty, total_marks, paper_name):
    raw_text = get_pdf_text(pdf_docs)
    if output_format == "MCQ":
        format_instruction = f"Generate {num_questions} multiple choice questions (MCQ) with 4 options each (A, B, C, D). For each question, list the options clearly and indicate the correct answer at the end in the format: 'Answer: <option letter>'. "
    elif output_format == "Short":
        format_instruction = f"Generate {num_questions} short answer questions."
    elif output_format == "Long":
        format_instruction = f"Generate {num_questions} long answer questions."
    else:
        format_instruction = f"Generate {num_questions} questions."
    prompt = f"""
    You are an expert exam paper generator. {format_instruction}
    Difficulty: {difficulty if difficulty != 'Any' else 'Mixed'}
    Total Marks/Duration: {total_marks if total_marks else 'Not specified'}
    Paper Name: {paper_name}
    
    Document Content:
    {raw_text}
    
    Output only the questions, clearly numbered. For MCQ, provide options A, B, C, D for each question, and indicate the correct answer as 'Answer: <option letter>'.
    """
    model = ChatGoogleGenerativeAI(model="gemini-2.0-flash", temperature=0.3)
    response = model.invoke(prompt)
    if hasattr(response, 'content'):
        response_text = response.content
    else:
        response_text = str(response)
    # For MCQ, group question, options, and answer together
    if output_format == "MCQ":
        questions = []
        lines = response_text.split('\n')
        current_q = []
        for line in lines:
            if line.strip().startswith(('Q', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) and current_q:
                questions.append('\n'.join(current_q).strip())
                current_q = [line]
            else:
                current_q.append(line)
        if current_q:
            questions.append('\n'.join(current_q).strip())
    else:
        questions = [q.strip() for q in response_text.split('\n') if q.strip() and (q.strip().startswith('Q') or q.strip()[0].isdigit())]
        if not questions:
            questions = response_text.split('\n')
    return questions



def main():
    st.set_page_config("Chat PDF")
    st.header("Chat with PDF using Gemini💁")

    with st.sidebar:
        st.title("Menu:")
        menu_option = st.radio("Choose an option:", ("Chat with PDF", "Generate Question Paper"))
        pdf_docs = st.file_uploader("Upload your PDF Files", accept_multiple_files=True)
        if menu_option == "Chat with PDF":
            if st.button("Submit & Process"):
                with st.spinner("Processing..."):
                    raw_text = get_pdf_text(pdf_docs)
                    text_chunks = get_text_chunks(raw_text)
                    get_vector_store(text_chunks)
                    st.success("Done")
        elif menu_option == "Generate Question Paper":
            if st.button("Submit & Process"):
                with st.spinner("Processing..."):
                    raw_text = get_pdf_text(pdf_docs)
                    text_chunks = get_text_chunks(raw_text)
                    get_vector_store(text_chunks)
                    st.success("Done")

    if menu_option == "Chat with PDF":
        user_question = st.text_input("Ask a Question from the PDF Files")
        if user_question:
            user_input(user_question)

    elif menu_option == "Generate Question Paper":
        st.subheader("Generate Question Paper from PDF")
        preview_data = None
        with st.form("question_paper_form"):
            school_name = st.text_input("School Name", value=st.session_state.get('reset_school_name', ''))
            exam_name = st.text_input("Exam Name", value=st.session_state.get('reset_exam_name', ''))
            class_name = st.text_input("Class", value=st.session_state.get('reset_class_name', ''))
            subject = st.text_input("Subject", value=st.session_state.get('reset_subject', ''))
            time = st.text_input("Time", value=st.session_state.get('reset_time', ''))
            max_marks = st.text_input("Maximum Marks (M.M.)", value=st.session_state.get('reset_max_marks', ''))
            notes_input = st.text_area("Notes (one per line)", value=st.session_state.get('reset_notes_input', ''))
            notes = notes_input.splitlines() if notes_input else []
            num_sections = st.number_input("Number of Sections", min_value=1, max_value=10, value=st.session_state.get('reset_num_sections', 2), key='num_sections')
            section_data = []
            for i in range(num_sections):
                st.markdown(f"---\n#### Section {chr(65+i)}")
                section_name = st.text_input(f"Section {chr(65+i)} Name", value=st.session_state.get(f'reset_section_name_{i}', f"Section {chr(65+i)}"), key=f"section_name_{i}")
                section_marks = st.text_input(f"Section {chr(65+i)} Marks", value=st.session_state.get(f'reset_section_marks_{i}', ''), key=f"section_marks_{i}")
                num_questions = st.number_input(f"Number of Questions in Section {chr(65+i)}", min_value=1, max_value=20, value=st.session_state.get(f'reset_num_questions_{i}', 2), key=f"num_questions_{i}")
                q_type = st.selectbox(f"Question Type for Section {chr(65+i)}", ["MCQ", "Short", "Long"], index=["MCQ", "Short", "Long"].index(st.session_state.get(f'reset_qtype_{i}', "MCQ")), key=f"qtype_{i}")
                q_marks = st.text_input(f"Marks per Question in Section {chr(65+i)}", value=st.session_state.get(f'reset_qmarks_{i}', ''), key=f"qmarks_{i}")
                section_data.append({
                    "section_name": section_name,
                    "section_marks": section_marks,
                    "num_questions": int(num_questions),
                    "q_type": q_type,
                    "q_marks": q_marks,
                    "questions": []
                })
            col1, col2 = st.columns([1,1])
            with col1:
                generate_btn = st.form_submit_button("Generate Paper")
            with col2:
                reset_btn = st.form_submit_button("Reset")
        if reset_btn:
            for k in list(st.session_state.keys()):
                if k.startswith('reset_') or k.startswith('section_name_') or k.startswith('section_marks_') or k.startswith('num_questions_') or k.startswith('qtype_') or k.startswith('qmarks_') or k == 'num_sections':
                    del st.session_state[k]
            st.session_state['show_download_buttons'] = False
            st.experimental_rerun()
        if generate_btn and pdf_docs:
            st.session_state['show_download_buttons'] = True
            st.session_state['last_qp_args'] = {
                'exam_name': exam_name,
                'section_data': section_data,
                'max_marks': max_marks,
                'school_name': school_name,
                'class_name': class_name,
                'subject': subject,
                'time': time,
                'notes': notes,
                'pdf_docs': pdf_docs
            }
            # Generate and save PDFs and DOCX only once per generation
            import tempfile, uuid, os
            unique_id = str(uuid.uuid4())
            base_dir = os.path.join(tempfile.gettempdir(), 'papermaker')
            os.makedirs(base_dir, exist_ok=True)
            pdf_path = os.path.join(base_dir, f"{unique_id}_question_paper.pdf")
            ans_path = os.path.join(base_dir, f"{unique_id}_answer_sheet.pdf")
            docx_path = os.path.join(base_dir, f"{unique_id}_question_paper.docx")
            # Generate questions for each section
            for section in section_data:
                questions = generate_questions_from_pdf(
                    pdf_docs,
                    section["num_questions"],
                    section["q_type"],
                    '', '', ''
                )
                section["questions"] = [
                    {"text": q, "marks": section["q_marks"], "type": section["q_type"]}
                    for q in questions
                ]
            # Generate files
            create_question_paper_pdf(
                paper_name=exam_name,
                questions_by_section=section_data,
                total_marks=max_marks,
                output_path=pdf_path,
                school_name=school_name,
                exam_name=exam_name,
                class_name=class_name,
                subject=subject,
                time=time,
                notes=notes,
                max_marks=max_marks
            )
            create_answer_sheet_pdf(
                paper_name=exam_name,
                questions_by_section=section_data,
                output_path=ans_path,
                school_name=school_name,
                exam_name=exam_name,
                class_name=class_name,
                subject=subject,
                time=time,
                max_marks=max_marks
            )
            import docx
            def create_question_paper_docx(paper_name, questions_by_section, total_marks, output_path, school_name, exam_name, class_name, subject, time, notes, max_marks):
                doc = docx.Document()
                doc.add_heading(school_name, 0)
                doc.add_paragraph(f"Exam: {exam_name}")
                doc.add_paragraph(f"Class: {class_name}")
                doc.add_paragraph(f"Subject: {subject}")
                doc.add_paragraph(f"Time: {time}")
                doc.add_paragraph(f"Maximum Marks: {max_marks}")
                if notes:
                    doc.add_paragraph("Notes:")
                    for note in notes:
                        doc.add_paragraph(note, style='List Bullet')
                for section in questions_by_section:
                    doc.add_heading(section['section_name'], level=1)
                    for idx, q in enumerate(section['questions']):
                        para = doc.add_paragraph()
                        para.add_run(f"Q{idx+1}. {q['text']}").bold = False
                        para.add_run(f"   [{q['marks']} marks]").italic = True
                doc.save(output_path)
                return output_path
            create_question_paper_docx(
                paper_name=exam_name,
                questions_by_section=section_data,
                total_marks=max_marks,
                output_path=docx_path,
                school_name=school_name,
                exam_name=exam_name,
                class_name=class_name,
                subject=subject,
                time=time,
                notes=notes,
                max_marks=max_marks
            )
            st.session_state['qp_pdf_path'] = pdf_path
            st.session_state['ans_pdf_path'] = ans_path
            st.session_state['preview_docx_path'] = docx_path
        # Use saved files for download/preview
        if st.session_state.get('show_download_buttons', False):
            pdf_path = st.session_state.get('qp_pdf_path')
            ans_path = st.session_state.get('ans_pdf_path')
            docx_path = st.session_state.get('preview_docx_path')
            st.markdown("#### Preview and Download")
            with open(pdf_path, "rb") as f:
                st.download_button("Download Question Paper", f, file_name=f"{st.session_state['last_qp_args']['exam_name'] or 'QuestionPaper'}.pdf", mime="application/pdf", key="download_qp_btn")
            with open(ans_path, "rb") as f:
                st.download_button("Download Answers", f, file_name=f"{st.session_state['last_qp_args']['exam_name'] or 'Answers'}.pdf", mime="application/pdf", key="download_ans_btn")
            if st.button("Reset", key="reset_btn"):
                st.session_state['show_download_buttons'] = False
                for k in ['qp_pdf_path', 'ans_pdf_path', 'preview_docx_path', 'last_qp_args']:
                    if k in st.session_state:
                        del st.session_state[k]


if __name__ == "__main__":
    main()