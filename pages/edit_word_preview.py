import streamlit as st
import tempfile
import uuid
from docx import Document
import os

try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

base_dir = os.path.join(tempfile.gettempdir(), 'papermaker')
os.makedirs(base_dir, exist_ok=True)

def main():
    st.title("Edit and Preview Question Paper (Word)")
    if 'preview_docx_path' in st.session_state:
        docx_path = st.session_state['preview_docx_path']
        # Load the entire document as a single editable block
        doc = Document(docx_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        edited_text = st.text_area("Edit the question paper (Word style)", value=full_text, height=600)
        if st.button("Save and Download as PDF"):
            unique_id = str(uuid.uuid4())
            edited_docx_path = os.path.join(base_dir, f"{unique_id}_edited_question_paper.docx")
            # Save the edited text as a new docx
            new_doc = Document()
            for para in edited_text.split("\n"):
                new_doc.add_paragraph(para)
            new_doc.save(edited_docx_path)
            # Convert to PDF if possible
            if DOCX2PDF_AVAILABLE:
                edited_pdf_path = os.path.join(base_dir, f"{unique_id}_edited_question_paper.pdf")
                try:
                    docx2pdf_convert(edited_docx_path, edited_pdf_path)
                    with open(edited_pdf_path, "rb") as f:
                        st.download_button("Download PDF", f, file_name="edited_question_paper.pdf", mime="application/pdf")
                except Exception as e:
                    st.error(f"PDF conversion failed: {e}")
                    with open(edited_docx_path, "rb") as f:
                        st.download_button("Download DOCX", f, file_name="edited_question_paper.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.warning("docx2pdf is not installed. Please install it for PDF export. Downloading as DOCX instead.")
                with open(edited_docx_path, "rb") as f:
                    st.download_button("Download DOCX", f, file_name="edited_question_paper.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        uploaded_docx = st.file_uploader("Upload the generated Word file to edit", type=["docx"])
        if uploaded_docx:
            unique_id = str(uuid.uuid4())
            uploaded_docx_path = os.path.join(base_dir, f"{unique_id}_uploaded_question_paper.docx")
            with open(uploaded_docx_path, "wb") as tmpfile:
                tmpfile.write(uploaded_docx.read())
            st.session_state['preview_docx_path'] = uploaded_docx_path
            st.rerun()

if __name__ == "__main__":
    main()
